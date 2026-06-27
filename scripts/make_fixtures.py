"""Generate the QA-03 test-document set in tests/fixtures.

Produces: a correct DOCX and PDF (Russian text), an empty PDF, a corrupt PDF,
and a wrong-type .txt file. Re-runnable; binaries live in tests/fixtures only.

Usage: python scripts/make_fixtures.py
Requires: python-docx, reportlab.
"""

from __future__ import annotations

import os

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

FX = os.path.join(os.path.dirname(__file__), "..", "tests", "fixtures")
UNICODE_FONT = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"


def make_docx(path: str) -> None:
    """Write a valid multi-paragraph Russian DOCX lecture."""
    doc = Document()
    doc.add_heading("Базы данных. Лекция 3", 0)
    for text in (
        "Реляционная модель данных основана на теории множеств и реляционной алгебре. " * 6,
        "Язык SQL используется для управления реляционными базами данных и выполнения запросов. " * 6,
        "Индексы ускоряют поиск данных в таблицах базы данных за счёт дополнительных структур. " * 6,
    ):
        doc.add_paragraph(text)
    doc.save(path)


def make_pdf(path: str) -> None:
    """Write a valid two-page Russian PDF lecture."""
    font = "Helvetica"
    try:
        pdfmetrics.registerFont(TTFont("ArialUni", UNICODE_FONT))
        font = "ArialUni"
    except Exception:  # noqa: BLE001 - fall back to a built-in font
        pass
    pdf = canvas.Canvas(path, pagesize=A4)
    lines = [
        "Алгоритмы и структуры данных. Лекция 5",
        "Сортировка слиянием имеет временную сложность O(n log n).",
        "Бинарный поиск работает за логарифмическое время O(log n) на отсортированном массиве.",
        "Хеш-таблицы обеспечивают доступ к данным в среднем за константное время O(1).",
        "Графовые алгоритмы: поиск в ширину BFS и поиск в глубину DFS обходят вершины графа.",
    ]
    y = 800
    for i, line in enumerate(lines):
        pdf.setFont(font, 14 if i == 0 else 11)
        pdf.drawString(50, y, line)
        y -= 30
    pdf.showPage()
    pdf.setFont(font, 11)
    pdf.drawString(50, 800, "Страница 2: динамическое программирование разбивает задачу на подзадачи.")
    pdf.showPage()
    pdf.save()


def main() -> None:
    os.makedirs(FX, exist_ok=True)
    make_docx(os.path.join(FX, "lecture_databases.docx"))
    make_pdf(os.path.join(FX, "lecture_algorithms.pdf"))
    open(os.path.join(FX, "empty.pdf"), "wb").close()
    with open(os.path.join(FX, "broken.pdf"), "wb") as fh:
        fh.write(b"%PDF-1.4 this is not a real pdf body \x00\x01 garbage")
    with open(os.path.join(FX, "notes.txt"), "w", encoding="utf-8") as fh:
        fh.write("plain text — wrong type, must be rejected")
    print("fixtures written to", os.path.abspath(FX))


if __name__ == "__main__":
    main()
