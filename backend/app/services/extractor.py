"""Text extraction from uploaded PDF and DOCX files."""

import os

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _resolve_kind(path: str, content_type: str | None) -> str:
    """Determine ``'pdf'`` or ``'docx'`` from content type or file extension."""
    ct = (content_type or "").lower()
    if ct == PDF_CONTENT_TYPE:
        return "pdf"
    if ct == DOCX_CONTENT_TYPE:
        return "docx"
    ext = path.rsplit(".", 1)[-1].lower() if "." in path else ""
    if ext in {"pdf", "docx"}:
        return ext
    raise ValueError(f"Unsupported file type for extraction: '{content_type or ext}'.")


def _extract_pdf(path: str) -> list[tuple[int, str]]:
    """Extract text per page from a PDF using pdfplumber."""
    import pdfplumber

    pages: list[tuple[int, str]] = []
    try:
        with pdfplumber.open(path) as pdf:
            for number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append((number, text))
    except ValueError:
        raise
    except Exception as exc:  # noqa: BLE001 - normalise library errors
        raise ValueError(f"Failed to read PDF: {exc}") from exc
    return pages


def _extract_docx(path: str) -> list[tuple[int, str]]:
    """Extract text from a DOCX as a single page (page 1)."""
    from docx import Document as DocxDocument

    try:
        document = DocxDocument(path)
        paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    paragraphs.append(" ".join(cells))
    except Exception as exc:  # noqa: BLE001 - normalise library errors
        raise ValueError(f"Failed to read DOCX: {exc}") from exc
    return [(1, "\n".join(paragraphs))]


def extract_text(path: str, content_type: str | None) -> list[tuple[int, str]]:
    """Extract text from a document file.

    Args:
        path: Absolute path to the stored file.
        content_type: Reported MIME type; the file extension is used as a
            fallback when the type is missing or generic.

    Returns:
        A list of ``(page_number, text)`` tuples. PDFs yield one entry per
        page; DOCX files yield a single entry numbered page 1.

    Raises:
        ValueError: If the file is missing, the type is unsupported, the file
            cannot be parsed, or no extractable text was found.
    """
    if not os.path.exists(path):
        raise ValueError(f"File not found: {path}")
    if os.path.getsize(path) == 0:
        raise ValueError("File is empty.")

    kind = _resolve_kind(path, content_type)
    pages = _extract_pdf(path) if kind == "pdf" else _extract_docx(path)

    if not any((text or "").strip() for _, text in pages):
        raise ValueError("No extractable text found in document.")

    return pages
